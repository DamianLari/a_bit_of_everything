import os
from docx import Document
from pylatex import Document as LaTeXDocument, Section, Subsection, Subsubsection, Command, Figure, Itemize, Enumerate
from pylatex.utils import NoEscape

def process_paragraph(paragraph):
    text = paragraph.text.strip()
    if not text:
        return None
    if paragraph.style.name.startswith('Heading 1'):
        return 'section', text
    elif paragraph.style.name.startswith('Heading 2'):
        return 'subsection', text
    elif paragraph.style.name.startswith('Heading 3'):
        return 'subsubsection', text
    elif paragraph.style.name.startswith('Heading 4'):
        return 'paragraph', text
    elif paragraph.style.name.startswith('Heading 5'):
        return 'subparagraph', text
    elif paragraph.style.name == 'Normal':
        return 'text', text
    elif paragraph.style.name in ['List Bullet', 'List Bullet 2', 'List Bullet 3']:
        return 'itemize', text
    elif paragraph.style.name in ['List Number', 'List Number 2', 'List Number 3']:
        return 'enumerate', text
    else:
        return 'text', text

def process_images(doc, image_folder):
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img = rel.target_part
            img_filename = os.path.basename(img.partname)
            img_data = img.blob
            img_path = os.path.join(image_folder, img_filename)
            with open(img_path, 'wb') as f:
                f.write(img_data)
            images.append(img_filename)
    return images

def convert_docx_to_latex(docx_file, output_file, image_folder="source"):
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)

    doc = Document(docx_file)
    latex_doc = LaTeXDocument()

    current_section = None
    current_subsection = None
    current_subsubsection = None
    current_paragraph = None
    current_list = None

    images = process_images(doc, image_folder)
    image_index = 0

    for paragraph in doc.paragraphs:
        result = process_paragraph(paragraph)
        if result:
            kind, content = result
            if kind == 'section':
                current_section = Section(content)
                latex_doc.append(current_section)
            elif kind == 'subsection':
                current_subsection = Subsection(content)
                current_section.append(current_subsection)
            elif kind == 'subsubsection':
                current_subsubsection = Subsubsection(content)
                current_subsection.append(current_subsubsection)
            elif kind == 'paragraph':
                current_paragraph = Command('paragraph', content)
                if current_subsubsection:
                    current_subsubsection.append(current_paragraph)
                elif current_subsection:
                    current_subsection.append(current_paragraph)
                elif current_section:
                    current_section.append(current_paragraph)
                else:
                    latex_doc.append(current_paragraph)
            elif kind == 'subparagraph':
                current_subparagraph = Command('subparagraph', content)
                if current_paragraph:
                    current_paragraph.append(current_subparagraph)
                elif current_subsubsection:
                    current_subsubsection.append(current_subparagraph)
                elif current_subsection:
                    current_subsection.append(current_subparagraph)
                elif current_section:
                    current_section.append(current_subparagraph)
                else:
                    latex_doc.append(current_subparagraph)
            elif kind == 'itemize':
                if current_list is None:
                    current_list = Itemize()
                    if current_subsubsection:
                        current_subsubsection.append(current_list)
                    elif current_subsection:
                        current_subsection.append(current_list)
                    elif current_section:
                        current_section.append(current_list)
                    else:
                        latex_doc.append(current_list)
                current_list.add_item(content)
            elif kind == 'enumerate':
                if current_list is None:
                    current_list = Enumerate()
                    if current_subsubsection:
                        current_subsubsection.append(current_list)
                    elif current_subsection:
                        current_subsection.append(current_list)
                    elif current_section:
                        current_section.append(current_list)
                    else:
                        latex_doc.append(current_list)
                current_list.add_item(content)
            elif kind == 'text':
                if current_list:
                    current_list = None  
                if current_subsubsection:
                    current_subsubsection.append(content)
                elif current_subsection:
                    current_subsection.append(content)
                elif current_section:
                    current_section.append(content)
                else:
                    latex_doc.append(content)

        for img_filename in images:
            with latex_doc.create(Figure(position='h!')) as figure:
                figure.add_image(os.path.join(image_folder, img_filename), width=NoEscape(r'0.8\textwidth'))
                figure.add_caption(f'Image {image_index + 1}')
            image_index += 1
    
    latex_doc.generate_tex(output_file)

# Usage
convert_docx_to_latex('input/InputFile.docx', 'OutputFile')
